from pathlib import Path
from docx import Document
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE_MD = BASE_DIR / "Project_Report_Assignment_2.md"
OUTPUT_DOCX = BASE_DIR / "Project_Report_Assignment_2.docx"


def build_docx_from_markdown():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source file not found: {SOURCE_MD}")

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    for line in lines:
        text = line.rstrip()

        if not text.strip():
            doc.add_paragraph("")
            continue

        if text.strip() == "---":
            doc.add_page_break()
            continue

        if text.startswith("# "):
            doc.add_heading(text[2:].strip(), level=1)
            continue

        if text.startswith("## "):
            doc.add_heading(text[3:].strip(), level=2)
            continue

        if text.startswith("### "):
            doc.add_heading(text[4:].strip(), level=3)
            continue

        if text.startswith("- "):
            doc.add_paragraph(text[2:].strip(), style="List Bullet")
            continue

        # Numbered list lines like "1. Item"
        if len(text) > 3 and text[0].isdigit() and text[1:3] == ". ":
            doc.add_paragraph(text[3:].strip(), style="List Number")
            continue

        doc.add_paragraph(text)

    doc.save(OUTPUT_DOCX)
    print(f"Created: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_docx_from_markdown()
